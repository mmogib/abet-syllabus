"""Generate filled ABET syllabus DOCX from template and SyllabusData."""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .assembler import SyllabusData


def generate_docx(
    data: SyllabusData,
    template_path: str | Path,
    output_path: str | Path,
) -> Path:
    """Fill the ABET template with course data and save as DOCX.

    Template structure (8 tables):
      Table 0: Title "COURSE SYLLABUS"
      Table 1: Course identification (dept, number, title)
      Table 2: Credit hours and categorization
      Table 3: Instructor
      Table 4: Textbooks
      Table 5: Specific course info (catalog desc, prereqs, designation)
      Table 6: CLO-SO mapping
      Table 7: Course topics

    Args:
        data: Assembled syllabus data.
        template_path: Path to the DOCX template.
        output_path: Where to save the filled DOCX.

    Returns:
        Path to the saved DOCX file.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    tables = doc.tables

    if len(tables) < 8:
        raise ValueError(
            f"Template has {len(tables)} tables, expected at least 8. "
            "Is this the correct ABET syllabus template?"
        )

    _fill_course_identification(tables[1], data)
    _fill_credits(tables[2], data)
    _fill_instructor(tables[3], data)
    _fill_textbooks(tables[4], data)
    _fill_specific_info(tables[5], data)
    _fill_clo_so_table(tables[6], data, doc)
    _fill_topics(tables[7], data)

    # Assessment data: the template has no assessment table, so we add one
    # after the topics table when assessment data is available.
    if data.assessments:
        _add_assessments_table(doc, data)

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Table fillers
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str) -> None:
    """Set cell text while preserving formatting of the first run if possible."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return

    p = paragraphs[0]
    if p.runs:
        # Preserve formatting from the first run
        first_run = p.runs[0]
        # Clear all runs
        for run in p.runs:
            run.text = ""
        first_run.text = text
    else:
        p.text = text

    # Remove extra paragraphs beyond the first
    for extra_p in paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)


def _fill_course_identification(table, data: SyllabusData) -> None:
    """Table 1: Department, Course Number, Course Title.

    Row 0: header (merged)
    Row 1: Department: | <value>
    Row 2: Course Number: | <value>
    Row 3: Course Title: | <value>
    """
    _set_cell_text(table.rows[1].cells[1], data.department)
    _set_cell_text(table.rows[2].cells[1], data.course_code)
    _set_cell_text(table.rows[3].cells[1], data.course_title)


def _fill_credits(table, data: SyllabusData) -> None:
    """Table 2: Credit hours and categorization.

    Row 0: header
    Row 1: "Credits, Contact Hrs" | <merged across 3 cols>
    Row 2: "Credits, Contact Hrs" | <merged, L/LAB/CR details>
    Row 3: "Credits Categorization" | Math and Basic Sciences | Engineering Topics | Other
    Row 4: "Credits Categorization" | <value> | <value> | <value>
    """
    # Row 1: lecture info
    lec_per_week = data.lecture_credits
    contact_text = f"{lec_per_week} Lecture/week (50 minutes)"
    if data.lab_credits > 0:
        contact_text += f", {data.lab_credits} Lab/week"
    _set_cell_text(table.rows[1].cells[1], contact_text)

    # Row 2: L, LAB, CR breakdown
    credit_detail = f"(L = {data.lecture_credits}, LAB = {data.lab_credits}, CR = {data.total_credits})"
    _set_cell_text(table.rows[2].cells[1], credit_detail)

    # Row 4: credit categorization values
    cats = data.credit_categories
    math_sci = cats.get("math_science", 0)
    eng = cats.get("engineering_cs", 0)
    other_val = cats.get("other", 0)
    # Also aggregate other categories into "other"
    other_total = (
        other_val
        + cats.get("humanities", 0)
        + cats.get("social_sciences_business", 0)
        + cats.get("general_education", 0)
    )

    if cats:
        _set_cell_text(table.rows[4].cells[1], str(math_sci) if math_sci else "")
        _set_cell_text(table.rows[4].cells[2], str(eng) if eng else "")
        _set_cell_text(table.rows[4].cells[3], str(other_total) if other_total else "")


def _fill_instructor(table, data: SyllabusData) -> None:
    """Table 3: Instructor name.

    Row 0: header
    Row 1: "Name" | <value>
    """
    _set_cell_text(table.rows[1].cells[1], data.instructor_name)


def _fill_textbooks(table, data: SyllabusData) -> None:
    """Table 4: Textbooks and supplemental materials.

    Row 0: header
    Row 1: "Textbook..." | <textbook text>
    Row 2: "Other Supplemental..." | <references>

    Groups textbooks by type: "required" goes in row 1, everything else
    (reference, recommended, electronic) goes in row 2.
    """
    if data.textbooks:
        required_texts = []
        supplemental_texts = []
        for tb in data.textbooks:
            if tb.textbook_type == "required":
                required_texts.append(tb.text)
            else:
                supplemental_texts.append(tb.text)

        _set_cell_text(table.rows[1].cells[1], "\n".join(required_texts))
        _set_cell_text(table.rows[2].cells[1], "\n".join(supplemental_texts))
    else:
        _set_cell_text(table.rows[1].cells[1], "")
        _set_cell_text(table.rows[2].cells[1], "")


def _fill_specific_info(table, data: SyllabusData) -> None:
    """Table 5: Catalog description, prerequisites, designation.

    Row 0: header
    Row 1: 5.a | "Course Content (Catalog Description)" | <description merged 3 cols>
    Row 2: 5.b | "Prerequisites or Co-requisites" | <prereqs merged>
    Row 3: 5.c | "Designation" | Required | Selected Elective | Elective
    Row 4: 5.c | "Designation" | <checkmark in appropriate column>
    """
    # Catalog description (row 1, cell 2 is merged across remaining cols)
    _set_cell_text(table.rows[1].cells[2], data.catalog_description)

    # Prerequisites and corequisites (row 2)
    prereq_text = data.prerequisites
    if data.corequisites:
        prereq_text += f"\nCo-requisites: {data.corequisites}"
    _set_cell_text(table.rows[2].cells[2], prereq_text)

    # Designation checkmarks (row 4)
    course_type_lower = data.course_type.lower()
    is_required = "required" in course_type_lower and "elective" not in course_type_lower
    is_selected_elective = "selected" in course_type_lower and "elective" in course_type_lower
    is_elective = "elective" in course_type_lower and "selected" not in course_type_lower

    _set_cell_text(table.rows[4].cells[2], "\u221A" if is_required else "")
    _set_cell_text(table.rows[4].cells[3], "\u221A" if is_selected_elective else "")
    _set_cell_text(table.rows[4].cells[4], "\u221A" if is_elective else "")


def _fill_clo_so_table(table, data: SyllabusData, doc: Document) -> None:
    """Table 6: CLO list with SO mappings.

    Template rows:
      Row 0: header "Specific Goals for the Course"
      Row 1: sub-header "6.a" | "6.a" | "6.b"
      Row 2: column headers "Specific Outcomes..." | (merged) | "Student Outcomes..."
      Row 3+: CLO rows: "CLO-1" | <text> | "SO-3"

    Strategy:
    - Keep rows 0-2 as-is (headers).
    - Use row 3 as template for data rows.
    - Remove existing data rows and add new ones.
    """
    if not data.clos:
        # Clear existing data rows if no CLOs
        _remove_table_rows(table, start_row=3)
        return

    # IMPORTANT: deepcopy the template row BEFORE removing it from the table.
    # Calling deepcopy on a detached XML element is fragile and may lose
    # parent-context-dependent attributes.
    template_row_element = None
    if len(table.rows) > 3:
        template_row_element = copy.deepcopy(table.rows[3]._tr)

    # Remove existing data rows (rows 3+)
    _remove_table_rows(table, start_row=3)

    # Add CLO rows
    for clo in data.clos:
        # Find SO mappings for this CLO
        so_text = ""
        if data.clo_so_matrix:
            for matrix_entry in data.clo_so_matrix:
                if matrix_entry["clo_label"] == clo.label:
                    mapped_sos = [
                        so_label
                        for so_label, mapped in matrix_entry["so_mappings"].items()
                        if mapped
                    ]
                    so_text = ", ".join(sorted(mapped_sos))
                    break

        if template_row_element is not None:
            new_row = copy.deepcopy(template_row_element)
            table._tbl.append(new_row)
            # The new row's cells
            row = table.rows[-1]
            _set_cell_text(row.cells[0], clo.label)
            _set_cell_text(row.cells[1], clo.text)
            _set_cell_text(row.cells[2], so_text)
        else:
            # Fallback: add a simple row
            row = table.add_row()
            row.cells[0].text = clo.label
            row.cells[1].text = clo.text
            if len(row.cells) > 2:
                row.cells[2].text = so_text


def _fill_topics(table, data: SyllabusData) -> None:
    """Table 7: Course topics.

    Row 0: header "Brief List of Course Topics (covered)"
    Row 1: topic text (single cell, all topics listed)

    The template has a single merged cell. We fill it with all topics.
    """
    if not data.topics:
        if len(table.rows) > 1:
            _set_cell_text(table.rows[1].cells[0], "")
        return

    # Format topics as a numbered list with hours
    lines = []
    for topic in data.topics:
        hours_str = f" ({topic.contact_hours:.0f}h)" if topic.contact_hours else ""
        lines.append(f"{topic.number}. {topic.title}{hours_str}")

    topic_text = "\n".join(lines)

    if len(table.rows) > 1:
        _set_cell_text(table.rows[1].cells[0], topic_text)
    else:
        row = table.add_row()
        row.cells[0].text = topic_text


def _add_assessments_table(doc: Document, data: SyllabusData) -> None:
    """Add an assessment methods table after the last table in the document.

    The ABET template does not include an assessment table, so we create one
    with columns: Assessment Task, Week Due, and Proportion (%).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree

    # Add a spacer paragraph before the table
    doc.add_paragraph()

    # Create the assessment table: header row + 1 row per assessment
    num_rows = 1 + len(data.assessments)
    table = doc.add_table(rows=num_rows, cols=3)

    # Apply borders via XML since 'Table Grid' style may not exist in template
    _apply_table_borders(table)

    # Header row
    headers = ["Assessment Task", "Week Due", "Proportion (%)"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, assessment in enumerate(data.assessments, start=1):
        row = table.rows[row_idx]
        row.cells[0].text = assessment.task
        row.cells[1].text = str(assessment.week_due) if assessment.week_due else ""
        row.cells[2].text = (
            f"{assessment.proportion:.0f}%"
            if assessment.proportion
            else ""
        )
        # Center-align week and proportion columns
        for col_idx in (1, 2):
            for paragraph in row.cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _apply_table_borders(table) -> None:
    """Apply single-line borders to a table via XML (style-independent)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = qn('w:tblPr')
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is not None:
        tbl_pr.remove(borders)

    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    from lxml import etree
    borders_element = etree.fromstring(borders_xml)
    tbl_pr.append(borders_element)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _remove_table_rows(table, start_row: int) -> None:
    """Remove rows from a table starting at start_row index."""
    tbl_element = table._tbl
    rows_to_remove = []
    for i, tr in enumerate(tbl_element.findall(qn('w:tr'))):
        if i >= start_row:
            rows_to_remove.append(tr)
    for tr in rows_to_remove:
        tbl_element.remove(tr)
