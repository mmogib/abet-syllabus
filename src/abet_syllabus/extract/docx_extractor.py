"""DOCX text and table extraction using python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree

from abet_syllabus.extract.models import ExtractedTable, ExtractionResult

# XML namespace for WordprocessingML
_WML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS = {"w": _WML_NS}


def _extract_cell_text(cell: Any) -> str:
    """Extract text from a table cell, including SDT (content control) text.

    python-docx's ``cell.text`` skips content inside ``<w:sdt>`` structured
    document tags (form-field content controls).  Many CRF2 DOCX files store
    field values (title, code, credits, etc.) inside SDTs, so ``cell.text``
    returns only the label.  This helper reads *all* ``<w:t>`` elements in the
    cell's XML, regardless of whether they are inside an SDT.
    """
    tc = cell._tc
    paragraphs: list[str] = []
    for p_elem in tc.findall(".//w:p", _NS):
        # Collect all text runs in this paragraph, including SDT content
        runs = p_elem.findall(".//w:r/w:t", _NS)
        p_text = "".join(r.text or "" for r in runs)
        paragraphs.append(p_text)
    return "\n".join(paragraphs).strip()


def extract_docx(file_path: str | Path) -> ExtractionResult:
    """Extract text and tables from a DOCX file.

    Uses python-docx to read all paragraphs and tables. Handles
    merged cells by reading the text content of each cell object.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        ExtractionResult with raw_text, tables, and metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file cannot be read as a DOCX.
    """
    path = Path(file_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Failed to read DOCX '{path.name}': {exc}") from exc

    # Extract paragraph text
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    raw_text = "\n".join(paragraphs)

    # Extract tables
    tables: list[ExtractedTable] = []
    for table_index, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        seen_row_keys: set[str] = set()

        for row in table.rows:
            cells = [_extract_cell_text(cell) for cell in row.cells]
            # Deduplicate rows that appear identical due to merged cells
            row_key = "\t".join(cells)
            if row_key in seen_row_keys:
                continue
            seen_row_keys.add(row_key)
            rows.append(cells)

        if not rows:
            continue

        header = rows[0] if rows else None
        tables.append(ExtractedTable(
            rows=rows,
            header=header,
            table_index=table_index,
        ))

    return ExtractionResult(
        raw_text=raw_text,
        tables=tables,
        file_path=str(path),
        file_extension=".docx",
        format_type="format_b_crf2",
        metadata={
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        },
    )
