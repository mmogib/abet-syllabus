"""Tests for ABET syllabus output generation (Stage 7)."""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from abet_syllabus.db.schema import init_db
from abet_syllabus.db import repository as repo
from abet_syllabus.db.models import (
    Course, CourseClo, CourseTextbook, CourseTopic,
    CourseAssessment, CreditCategorization, CourseInstructor,
    PloDefinition, Program,
)
from abet_syllabus.generate.assembler import (
    SyllabusData, SyllabusCLO, SyllabusTopic, SyllabusAssessment,
    SyllabusTextbook, assemble_syllabus_data, _renumber_clos, _build_clo_so_matrix,
)
from abet_syllabus.generate.generator import (
    GenerateResult, generate_syllabus, generate_program, _make_output_filename,
)
from abet_syllabus.generate.pdf_converter import is_pdf_available


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

def _find_template() -> Path | None:
    """Find the ABET template in the project tree."""
    # Try various locations
    candidates = [
        Path("resources/templates/ABETSyllabusTemplate.docx"),
        Path(__file__).resolve().parent.parent / "resources" / "templates" / "ABETSyllabusTemplate.docx",
        # Main repo (worktrees may not have resources/)
        Path("C:/Users/mmogi/Projects/published_apps/abet-syllabus/resources/templates/ABETSyllabusTemplate.docx"),
    ]
    for p in candidates:
        if p.exists():
            return p
    return None


TEMPLATE_PATH = _find_template()

# Skip tests that need the template if it's not available
needs_template = pytest.mark.skipif(
    TEMPLATE_PATH is None,
    reason="ABET template not found (resources not in worktree)"
)


@pytest.fixture
def db_path(tmp_path):
    """Create a temporary database with test data."""
    db_file = tmp_path / "test.db"
    conn = init_db(str(db_file))
    _populate_test_data(conn)
    conn.close()
    return db_file


@pytest.fixture
def db_conn():
    """Create an in-memory database with test data."""
    conn = init_db(":memory:")
    _populate_test_data(conn)
    yield conn
    conn.close()


@pytest.fixture
def empty_db_path(tmp_path):
    """Create a temporary empty database."""
    db_file = tmp_path / "empty.db"
    conn = init_db(str(db_file))
    conn.close()
    return db_file


def _populate_test_data(conn):
    """Insert test courses, CLOs, topics, textbooks, assessments."""
    # Program
    repo.upsert_program(conn, Program(program_code="MATH", program_name="Mathematics"))
    repo.upsert_program(conn, Program(program_code="DATA", program_name="Data Science"))

    # PLO definitions for MATH
    for i in range(1, 8):
        repo.upsert_plo(conn, PloDefinition(
            program_code="MATH",
            plo_code=f"MATH_PLO_{i}",
            plo_label=str(i),
            plo_description=f"PLO {i} description",
            sequence=i,
        ))

    # Course 1: MATH 101 (with full data)
    c1_id = repo.upsert_course(conn, Course(
        course_code="MATH 101",
        course_title="Calculus I",
        department="Mathematics",
        college="College of Sciences",
        catalog_description="Introduction to differential and integral calculus.",
        credit_hours_raw="3-0-3",
        lecture_credits=3,
        lab_credits=0,
        total_credits=3,
        course_type="Required",
        level="Undergraduate",
        prerequisites="MATH 001",
        corequisites="",
    ))
    repo.link_course_program(conn, c1_id, "MATH")

    # CLOs for MATH 101 with hierarchical codes
    clo_ids = repo.replace_course_clos(conn, c1_id, [
        CourseClo(course_id=c1_id, clo_code="1.1", clo_category="Knowledge",
                  clo_text="Define limits and continuity.", sequence=1),
        CourseClo(course_id=c1_id, clo_code="1.2", clo_category="Knowledge",
                  clo_text="Describe fundamental theorem of calculus.", sequence=2),
        CourseClo(course_id=c1_id, clo_code="2.1", clo_category="Skills",
                  clo_text="Compute derivatives of standard functions.", sequence=3),
        CourseClo(course_id=c1_id, clo_code="2.2", clo_category="Skills",
                  clo_text="Evaluate definite and indefinite integrals.", sequence=4),
    ])

    # CLO-PLO mappings
    plos = repo.get_plos_for_program(conn, "MATH")
    # Map CLO 1.1 -> PLO 1, PLO 3
    repo.upsert_clo_plo_mapping(conn, __import__('abet_syllabus.db.models', fromlist=['CloPloMapping']).CloPloMapping(
        course_clo_id=clo_ids[0], plo_id=plos[0].id, program_code="MATH",
        mapping_source="test", confidence=1.0,
    ))
    repo.upsert_clo_plo_mapping(conn, __import__('abet_syllabus.db.models', fromlist=['CloPloMapping']).CloPloMapping(
        course_clo_id=clo_ids[0], plo_id=plos[2].id, program_code="MATH",
        mapping_source="test", confidence=1.0,
    ))
    # Map CLO 2.1 -> PLO 5
    repo.upsert_clo_plo_mapping(conn, __import__('abet_syllabus.db.models', fromlist=['CloPloMapping']).CloPloMapping(
        course_clo_id=clo_ids[2], plo_id=plos[4].id, program_code="MATH",
        mapping_source="test", confidence=1.0,
    ))

    # Topics
    repo.replace_course_topics(conn, c1_id, [
        CourseTopic(course_id=c1_id, topic_number=1, topic_title="Limits and Continuity",
                    contact_hours=6.0, sequence=1),
        CourseTopic(course_id=c1_id, topic_number=2, topic_title="Derivatives",
                    contact_hours=9.0, sequence=2),
        CourseTopic(course_id=c1_id, topic_number=3, topic_title="Applications of Derivatives",
                    contact_hours=6.0, sequence=3),
        CourseTopic(course_id=c1_id, topic_number=4, topic_title="Integration",
                    contact_hours=9.0, sequence=4),
    ])

    # Textbooks
    repo.replace_course_textbooks(conn, c1_id, [
        CourseTextbook(course_id=c1_id, textbook_text="Calculus by James Stewart, 8th ed.",
                       textbook_type="required", sequence=1),
        CourseTextbook(course_id=c1_id, textbook_text="Thomas' Calculus, 14th ed.",
                       textbook_type="reference", sequence=2),
    ])

    # Assessments
    repo.replace_course_assessment(conn, c1_id, [
        CourseAssessment(course_id=c1_id, assessment_task="Midterm Exam 1",
                         week_due="7", proportion=25.0, sequence=1),
        CourseAssessment(course_id=c1_id, assessment_task="Midterm Exam 2",
                         week_due="12", proportion=25.0, sequence=2),
        CourseAssessment(course_id=c1_id, assessment_task="Final Exam",
                         week_due="16", proportion=40.0, sequence=3),
        CourseAssessment(course_id=c1_id, assessment_task="Homework",
                         week_due="", proportion=10.0, sequence=4),
    ])

    # Credit categorization
    repo.upsert_credit_categorization(conn, CreditCategorization(
        course_id=c1_id, math_science=3.0, engineering_cs=0.0,
    ))

    # Instructor
    repo.upsert_instructor(conn, CourseInstructor(
        course_id=c1_id, instructor_name="Dr. Ahmed",
        term_code="T252", role="coordinator",
    ))

    # Course 2: MATH 201 (minimal data, no CLOs)
    c2_id = repo.upsert_course(conn, Course(
        course_code="MATH 201",
        course_title="Calculus III",
        department="Mathematics",
        college="College of Sciences",
        catalog_description="Multivariable calculus.",
        credit_hours_raw="3-0-3",
        lecture_credits=3,
        lab_credits=0,
        total_credits=3,
        course_type="Required",
    ))
    repo.link_course_program(conn, c2_id, "MATH")


# ---------------------------------------------------------------------------
# Tests: CLO renumbering
# ---------------------------------------------------------------------------

class TestCloRenumbering:
    def test_renumber_hierarchical_codes(self, db_conn):
        """Hierarchical codes like 1.1, 1.2, 2.1 become CLO-1, CLO-2, CLO-3."""
        course = repo.get_course(db_conn, "MATH 101")
        clos = repo.get_course_clos(db_conn, course.id)
        renumbered = _renumber_clos(clos)

        assert len(renumbered) == 4
        assert renumbered[0].label == "CLO-1"
        assert renumbered[1].label == "CLO-2"
        assert renumbered[2].label == "CLO-3"
        assert renumbered[3].label == "CLO-4"

    def test_renumber_preserves_text(self, db_conn):
        course = repo.get_course(db_conn, "MATH 101")
        clos = repo.get_course_clos(db_conn, course.id)
        renumbered = _renumber_clos(clos)

        assert renumbered[0].text == "Define limits and continuity."
        assert renumbered[0].original_code == "1.1"

    def test_renumber_empty_list(self):
        assert _renumber_clos([]) == []


# ---------------------------------------------------------------------------
# Tests: CLO-SO matrix
# ---------------------------------------------------------------------------

class TestCloSoMatrix:
    def test_build_matrix_with_mappings(self, db_conn):
        course = repo.get_course(db_conn, "MATH 101")
        clos = repo.get_course_clos(db_conn, course.id)
        renumbered = _renumber_clos(clos)
        plos = repo.get_plos_for_program(db_conn, "MATH")
        plo_labels = [p.plo_label for p in plos]
        mappings = repo.get_mappings_for_course(db_conn, course.id, "MATH")

        matrix = _build_clo_so_matrix(renumbered, mappings, plo_labels)

        assert len(matrix) == 4
        # CLO-1 maps to SO-1, SO-3
        assert matrix[0]["clo_label"] == "CLO-1"
        assert matrix[0]["so_mappings"]["SO-1"] is True
        assert matrix[0]["so_mappings"]["SO-3"] is True
        assert matrix[0]["so_mappings"]["SO-2"] is False

        # CLO-3 maps to SO-5
        assert matrix[2]["clo_label"] == "CLO-3"
        assert matrix[2]["so_mappings"]["SO-5"] is True

    def test_build_matrix_empty(self):
        matrix = _build_clo_so_matrix([], [], [])
        assert matrix == []


# ---------------------------------------------------------------------------
# Tests: SyllabusData assembly
# ---------------------------------------------------------------------------

class TestAssembleSyllabusData:
    def test_assemble_full_course(self, db_path):
        data = assemble_syllabus_data(db_path, "MATH 101", program_code="MATH", term="T252")

        assert data.course_code == "MATH 101"
        assert data.course_title == "Calculus I"
        assert data.department == "Mathematics"
        assert data.total_credits == 3
        assert data.prerequisites == "MATH 001"
        assert data.course_type == "Required"
        assert data.instructor_name == "Dr. Ahmed"
        assert data.term_code == "T252"

        assert len(data.clos) == 4
        assert data.clos[0].label == "CLO-1"

        assert len(data.topics) == 4
        assert data.topics[0].title == "Limits and Continuity"

        assert len(data.textbooks) == 2
        assert isinstance(data.textbooks[0], SyllabusTextbook)
        assert len(data.assessments) == 4

        assert len(data.clo_so_matrix) == 4
        assert data.credit_categories.get("math_science") == 3.0

    def test_assemble_without_program(self, db_path):
        data = assemble_syllabus_data(db_path, "MATH 101")
        assert data.clo_so_matrix == []
        assert len(data.clos) == 4

    def test_assemble_minimal_course(self, db_path):
        data = assemble_syllabus_data(db_path, "MATH 201")
        assert data.course_code == "MATH 201"
        assert data.clos == []
        assert data.topics == []
        assert data.textbooks == []

    def test_assemble_course_not_found(self, db_path):
        with pytest.raises(ValueError, match="Course not found"):
            assemble_syllabus_data(db_path, "NONEXIST 999")

    def test_instructor_override(self, db_path):
        data = assemble_syllabus_data(
            db_path, "MATH 101", term="T252", instructor="Dr. Override"
        )
        assert data.instructor_name == "Dr. Override"


# ---------------------------------------------------------------------------
# Tests: Output filename
# ---------------------------------------------------------------------------

class TestOutputFilename:
    def test_with_term(self):
        assert _make_output_filename("MATH 101", "T252") == "T252_MATH_101_ABET_Syllabus"

    def test_without_term(self):
        assert _make_output_filename("MATH 101", None) == "MATH_101_ABET_Syllabus"

    def test_no_spaces_in_code(self):
        name = _make_output_filename("BUS200", "T252")
        assert " " not in name


# ---------------------------------------------------------------------------
# Tests: DOCX generation
# ---------------------------------------------------------------------------

class TestDocxGeneration:
    @needs_template
    def test_generate_docx_creates_file(self, db_path, tmp_path):
        data = assemble_syllabus_data(db_path, "MATH 101", program_code="MATH", term="T252")
        output = tmp_path / "test_output.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        result = generate_docx(data, TEMPLATE_PATH, output)

        assert result.exists()
        assert result.stat().st_size > 0

    @needs_template
    def test_generate_docx_valid_content(self, db_path, tmp_path):
        """Verify the generated DOCX has expected content in tables."""
        data = assemble_syllabus_data(db_path, "MATH 101", program_code="MATH", term="T252")
        output = tmp_path / "test_content.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        from docx import Document

        generate_docx(data, TEMPLATE_PATH, output)

        doc = Document(str(output))
        tables = doc.tables

        # Table 1: course identification
        assert "Mathematics" in tables[1].rows[1].cells[1].text
        assert "MATH 101" in tables[1].rows[2].cells[1].text
        assert "Calculus I" in tables[1].rows[3].cells[1].text

        # Table 3: instructor
        assert "Dr. Ahmed" in tables[3].rows[1].cells[1].text

        # Table 6: CLO-SO (should have CLO rows)
        assert len(tables[6].rows) >= 4  # header rows + at least 1 CLO

    @needs_template
    def test_generate_docx_minimal_data(self, db_path, tmp_path):
        """Course with no CLOs/topics should still generate a valid DOCX."""
        data = assemble_syllabus_data(db_path, "MATH 201")
        output = tmp_path / "test_minimal.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        result = generate_docx(data, TEMPLATE_PATH, output)

        assert result.exists()
        assert result.stat().st_size > 0

    def test_generate_docx_missing_template(self, db_path, tmp_path):
        data = assemble_syllabus_data(db_path, "MATH 101")
        output = tmp_path / "test.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        with pytest.raises(FileNotFoundError):
            generate_docx(data, "/nonexistent/template.docx", output)


# ---------------------------------------------------------------------------
# Tests: generate_syllabus (end-to-end)
# ---------------------------------------------------------------------------

class TestGenerateSyllabus:
    @needs_template
    def test_generate_single_course(self, db_path, tmp_path):
        result = generate_syllabus(
            db_path=db_path,
            course_code="MATH 101",
            program_code="MATH",
            term="T252",
            template_path=TEMPLATE_PATH,
            output_dir=tmp_path,
            pdf=False,
        )

        assert result.status == "success"
        assert result.docx_path is not None
        assert Path(result.docx_path).exists()
        assert "T252_MATH_101_ABET_Syllabus.docx" in result.docx_path

    @needs_template
    def test_generate_with_instructor(self, db_path, tmp_path):
        result = generate_syllabus(
            db_path=db_path,
            course_code="MATH 101",
            term="T252",
            instructor="Dr. Custom",
            template_path=TEMPLATE_PATH,
            output_dir=tmp_path,
            pdf=False,
        )

        assert result.status == "success"

    def test_generate_course_not_found(self, db_path, tmp_path):
        result = generate_syllabus(
            db_path=db_path,
            course_code="FAKE 999",
            template_path=TEMPLATE_PATH or "fake.docx",
            output_dir=tmp_path,
            pdf=False,
        )

        assert result.status == "error"
        assert "not found" in result.message.lower()


# ---------------------------------------------------------------------------
# Tests: generate_program
# ---------------------------------------------------------------------------

class TestGenerateProgram:
    @needs_template
    def test_generate_all_in_program(self, db_path, tmp_path):
        results = generate_program(
            db_path=db_path,
            program_code="MATH",
            term="T252",
            template_path=TEMPLATE_PATH,
            output_dir=tmp_path,
            pdf=False,
        )

        assert len(results) == 2  # MATH 101 and MATH 201
        successes = [r for r in results if r.status == "success"]
        assert len(successes) == 2

    def test_generate_program_not_found(self, db_path, tmp_path):
        results = generate_program(
            db_path=db_path,
            program_code="NONEXIST",
            template_path=TEMPLATE_PATH or "fake.docx",
            output_dir=tmp_path,
            pdf=False,
        )

        assert len(results) == 1
        assert results[0].status == "error"


# ---------------------------------------------------------------------------
# Tests: PDF conversion
# ---------------------------------------------------------------------------

class TestPdfConversion:
    def test_is_pdf_available_returns_bool(self):
        result = is_pdf_available()
        assert isinstance(result, bool)

    @needs_template
    @pytest.mark.skipif(not is_pdf_available(), reason="No PDF converter available")
    def test_pdf_conversion(self, db_path, tmp_path):
        from abet_syllabus.generate.pdf_converter import convert_to_pdf

        data = assemble_syllabus_data(db_path, "MATH 101")
        docx_out = tmp_path / "test.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        generate_docx(data, TEMPLATE_PATH, docx_out)

        pdf_out = tmp_path / "test.pdf"
        result = convert_to_pdf(docx_out, pdf_out)
        assert result.exists()

    def test_convert_nonexistent_file(self, tmp_path):
        from abet_syllabus.generate.pdf_converter import convert_to_pdf
        with pytest.raises(FileNotFoundError):
            convert_to_pdf(tmp_path / "nonexistent.docx")


# ---------------------------------------------------------------------------
# Tests: GenerateResult dataclass
# ---------------------------------------------------------------------------

class TestGenerateResult:
    def test_success_result(self):
        r = GenerateResult(
            course_code="MATH 101",
            docx_path="/path/to/file.docx",
            pdf_path=None,
            status="success",
            message="Generated",
        )
        assert r.status == "success"

    def test_error_result(self):
        r = GenerateResult(
            course_code="MATH 101",
            docx_path=None,
            pdf_path=None,
            status="error",
            message="Something failed",
        )
        assert r.status == "error"


# ---------------------------------------------------------------------------
# Tests: MAJOR-1 — Textbook type preservation
# ---------------------------------------------------------------------------

class TestTextbookTypePreservation:
    def test_textbook_types_preserved_in_assembly(self, db_path):
        """Textbook types from the DB must be preserved in SyllabusData."""
        data = assemble_syllabus_data(db_path, "MATH 101")
        assert len(data.textbooks) == 2

        required_tbs = [tb for tb in data.textbooks if tb.textbook_type == "required"]
        reference_tbs = [tb for tb in data.textbooks if tb.textbook_type == "reference"]

        assert len(required_tbs) == 1
        assert "Stewart" in required_tbs[0].text
        assert len(reference_tbs) == 1
        assert "Thomas" in reference_tbs[0].text

    @needs_template
    def test_textbooks_grouped_in_docx(self, db_path, tmp_path):
        """Required textbooks go in row 1, references in row 2 of table 4."""
        data = assemble_syllabus_data(db_path, "MATH 101")
        output = tmp_path / "test_textbooks.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        from docx import Document

        generate_docx(data, TEMPLATE_PATH, output)
        doc = Document(str(output))

        # Table 4: textbooks
        tb_table = doc.tables[4]
        required_cell = tb_table.rows[1].cells[1].text
        reference_cell = tb_table.rows[2].cells[1].text

        assert "Stewart" in required_cell
        assert "Thomas" in reference_cell


# ---------------------------------------------------------------------------
# Tests: MAJOR-2 — Safe deepcopy of CLO-SO template row
# ---------------------------------------------------------------------------

class TestCloSoDeepCopy:
    @needs_template
    def test_multiple_clos_render_correctly(self, db_path, tmp_path):
        """Multiple CLOs should each get their own row via deepcopy."""
        data = assemble_syllabus_data(db_path, "MATH 101", program_code="MATH")
        output = tmp_path / "test_clos.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        from docx import Document

        generate_docx(data, TEMPLATE_PATH, output)
        doc = Document(str(output))

        # Table 6: CLO-SO mapping. 3 header rows + 4 CLO rows = 7 rows
        clo_table = doc.tables[6]
        assert len(clo_table.rows) == 7  # 3 header + 4 CLOs

        # Verify each CLO row has correct label
        assert "CLO-1" in clo_table.rows[3].cells[0].text
        assert "CLO-2" in clo_table.rows[4].cells[0].text
        assert "CLO-3" in clo_table.rows[5].cells[0].text
        assert "CLO-4" in clo_table.rows[6].cells[0].text


# ---------------------------------------------------------------------------
# Tests: MAJOR-3 — Assessment table rendering
# ---------------------------------------------------------------------------

class TestAssessmentRendering:
    @needs_template
    def test_assessments_appear_in_generated_docx(self, db_path, tmp_path):
        """Assessment data should be rendered as a table in the output DOCX."""
        data = assemble_syllabus_data(db_path, "MATH 101")
        output = tmp_path / "test_assessments.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        from docx import Document

        generate_docx(data, TEMPLATE_PATH, output)
        doc = Document(str(output))

        # The template has 8 tables; assessments add a 9th
        assert len(doc.tables) == 9

        # The assessment table is the last one
        assess_table = doc.tables[8]
        # Header row + 4 assessment rows
        assert len(assess_table.rows) == 5

        # Verify header
        assert "Assessment Task" in assess_table.rows[0].cells[0].text
        assert "Week Due" in assess_table.rows[0].cells[1].text
        assert "Proportion" in assess_table.rows[0].cells[2].text

        # Verify data rows
        assert "Midterm Exam 1" in assess_table.rows[1].cells[0].text
        assert "25%" in assess_table.rows[1].cells[2].text
        assert "Final Exam" in assess_table.rows[3].cells[0].text
        assert "40%" in assess_table.rows[3].cells[2].text

    @needs_template
    def test_no_assessments_no_extra_table(self, db_path, tmp_path):
        """When there are no assessments, no extra table should be added."""
        data = assemble_syllabus_data(db_path, "MATH 201")
        output = tmp_path / "test_no_assessments.docx"

        from abet_syllabus.generate.docx_generator import generate_docx
        from docx import Document

        generate_docx(data, TEMPLATE_PATH, output)
        doc = Document(str(output))

        # Only the original 8 tables, no assessment table added
        assert len(doc.tables) == 8
